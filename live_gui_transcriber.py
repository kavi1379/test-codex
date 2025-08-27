import tkinter as tk
from tkinter import scrolledtext, filedialog, simpledialog, messagebox
import threading
import pyaudio
import wave
import whisper
from docx import Document
from docx.shared import Inches
import os
import time
import tempfile
import keyboard
import nltk
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import re
import queue
from PIL import ImageGrab, Image

nltk_path = os.path.abspath(os.path.join(os.path.dirname(__file__), 'nltk_data'))
nltk.data.path.append(nltk_path)

CHUNK = 2048
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 16000
RECORD_SECONDS = 5
FILTER_PHRASES = [
    "thanks for watching", "subscribe", "video", "click the bell",
    "follow me", "don't forget", "hit like", "see you next time"
]

def simple_sentence_tokenizer(text):
    tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')
    return tokenizer.tokenize(text)

class LiveTranscriberApp:
    def __init__(self, root, word_path, model_name="small"):
        self.root = root
        self.word_path = word_path
        self.model_name = model_name
        self.model = whisper.load_model(model_name)
        self.doc = Document(word_path)
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')

        self.p = pyaudio.PyAudio()
        self.stream = self.p.open(format=FORMAT, channels=CHANNELS, rate=RATE,
                                  input=True, frames_per_buffer=CHUNK)
        self.recording = True
        self.slide_id = 1
        self.topic_label = tk.StringVar(value="Topic: [Manual]")
        self.current_line = ""

        self.setup_ui()
        self.bind_hotkeys()
        self.audio_queue = queue.Queue()
        self.start_threads()

    def setup_ui(self):
        self.root.title("🎙️ Live Transcriber GUI")
        self.root.geometry("850x750")

        self.slide_label = tk.Label(self.root, text=f"Slide {self.slide_id}", font=("Arial", 16))
        self.slide_label.pack(pady=5)

        self.topic_display = tk.Label(self.root, textvariable=self.topic_label, font=("Arial", 12), fg="blue")
        self.topic_display.pack(pady=2)

        self.transcript_box = scrolledtext.ScrolledText(self.root, height=18, font=("Consolas", 12), undo=True, wrap=tk.WORD)
        self.transcript_box.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

        self.note_entry = tk.Entry(self.root, font=("Arial", 12), width=80, justify="right")
        self.note_entry.pack(pady=5)
        self.note_entry.bind('<Return>', lambda e: self.add_note())
        self.note_entry.bind('<Control-a>', lambda e: self.note_entry.select_range(0, tk.END))

        btn_frame = tk.Frame(self.root)
        btn_frame.pack(pady=10)

        tk.Button(btn_frame, text="✅ Save & Next", command=self.save_slide, width=15).grid(row=0, column=0, padx=5)
        tk.Button(btn_frame, text="🔙 Previous Slide", command=self.previous_slide, width=15).grid(row=0, column=1, padx=5)
        tk.Button(btn_frame, text="📝 Set Topic", command=self.set_topic, width=15).grid(row=0, column=2, padx=5)
        tk.Button(btn_frame, text="📋 Paste Image", command=self.paste_clipboard_image, width=15).grid(row=0, column=3, padx=5)
        tk.Button(btn_frame, text="💾 Save All", command=self.save_all_document, width=12).grid(row=0, column=4, padx=5)
        tk.Button(btn_frame, text="❌ Exit", command=self.exit_app, width=10).grid(row=0, column=5, padx=5)

    def bind_hotkeys(self):
        keyboard.add_hotkey('alt+s', self.save_slide)
        keyboard.add_hotkey('alt+q', lambda: self.root.after(0, self.exit_app))
        keyboard.add_hotkey('alt+p', self.previous_slide)
        keyboard.add_hotkey('ctrl+v', self.paste_clipboard_image)

    def clean_text(self, text):
        lowered = text.lower()
        for phrase in FILTER_PHRASES:
            if phrase in lowered:
                return ""
        return text.strip()

    def start_threads(self):
        threading.Thread(target=self.recorder_loop, daemon=True).start()
        threading.Thread(target=self.transcriber_loop, daemon=True).start()

    def recorder_loop(self):
        while self.recording:
            frames = []
            for _ in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
                data = self.stream.read(CHUNK, exception_on_overflow=False)
                frames.append(data)
            self.audio_queue.put(frames)

    def transcriber_loop(self):
        while self.recording:
            if not self.audio_queue.empty():
                frames = self.audio_queue.get()
                filename = os.path.join(tempfile.gettempdir(), f"slide_{self.slide_id}_{int(time.time())}.wav")
                with wave.open(filename, 'wb') as wf:
                    wf.setnchannels(CHANNELS)
                    wf.setsampwidth(self.p.get_sample_size(FORMAT))
                    wf.setframerate(RATE)
                    wf.writeframes(b''.join(frames))

                try:
                    result = self.model.transcribe(filename, language="en", fp16=False)
                    cleaned = self.clean_text(result['text'])
                    if not cleaned:
                        continue
                    for s in simple_sentence_tokenizer(cleaned):
                        self.transcript_box.insert(tk.END, s.strip() + "\n")
                        self.transcript_box.see(tk.END)
                except Exception as e:
                    print(f"Transcribe error: {e}")

    def add_note(self):
        note = self.note_entry.get().strip()
        if note:
            self.transcript_box.insert(tk.END, f"[Note]: {note}\n")
            self.transcript_box.see(tk.END)
            self.note_entry.delete(0, tk.END)

    def save_slide(self):
        transcript = self.transcript_box.get(1.0, tk.END).strip()
        if transcript:
            self.doc.add_paragraph(f"Transcript for Slide {self.slide_id} ({self.topic_label.get()}):")
            self.doc.add_paragraph(transcript)
            self.doc.add_paragraph("\n")
            self.doc.save(self.word_path)

        self.slide_id += 1
        self.slide_label.config(text=f"Slide {self.slide_id}")
        self.transcript_box.delete(1.0, tk.END)

    def save_all_document(self):
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if path:
            self.doc.save(path)

    def previous_slide(self):
        if self.slide_id > 1:
            self.slide_id -= 1
            self.slide_label.config(text=f"Slide {self.slide_id}")

    def set_topic(self):
        topic = simpledialog.askstring("Set Topic", "Enter slide topic:")
        if topic:
            self.topic_label.set(f"Topic: {topic}")

    def paste_clipboard_image(self):
        try:
            image = ImageGrab.grabclipboard()
            if isinstance(image, Image.Image):
                temp_path = os.path.join(tempfile.gettempdir(), f"pasted_{int(time.time())}.png")
                image.save(temp_path)
                self.doc.add_picture(temp_path, width=Inches(5.5))
                self.doc.add_paragraph("Pasted image added.")
                self.doc.save(self.word_path)
            else:
                messagebox.showwarning("No Image", "Clipboard does not contain an image.")
        except Exception as e:
            messagebox.showerror("Paste Error", str(e))

    def exit_app(self):
        self.recording = False
        self.stream.stop_stream()
        self.stream.close()
        self.p.terminate()
        self.root.quit()

if __name__ == "__main__":
    from tkinter import Toplevel, Label, Button, StringVar
    from tkinter import ttk

    root = tk.Tk()
    root.withdraw()

    word_path = filedialog.askopenfilename(
        title="Select Word File",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not word_path:
        print("❌ No file selected.")
        exit()

    def choose_model():
        model_window = Toplevel(root)
        model_window.title("Choose Whisper Model")
        model_window.geometry("300x150")

        var = StringVar(model_window)
        var.set("small")

        Label(model_window, text="Select Whisper model:").pack(pady=10)
        dropdown = ttk.Combobox(model_window, textvariable=var, values=["tiny", "small", "medium", "large"], state="readonly")
        dropdown.pack(pady=5)

        def submit():
            model_window.model = var.get()
            model_window.destroy()

        Button(model_window, text="OK", command=submit).pack(pady=10)
        root.wait_window(model_window)
        return getattr(model_window, 'model', None)

    model_name = choose_model()
    if not model_name:
        print("❌ No model selected.")
        exit()

    root.deiconify()
    app = LiveTranscriberApp(root, word_path, model_name=model_name.lower())
    root.mainloop()
